from fastapi import APIRouter, UploadFile, File, HTTPException, BackgroundTasks
from fastapi.responses import FileResponse
from pdf2image import convert_from_path
import pytesseract
from docx import Document
from docx.shared import Inches
import os, uuid, time

router = APIRouter()

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ================= IMAGE TO EDITABLE WORD LOGIC =================
def pdf_to_images_to_word(pdf_path, docx_path):
    # 1. PDF ni High Quality Images (DPI 300) ga marusthunnam
    images = convert_from_path(pdf_path, dpi=300)
    
    word_doc = Document()
    
    for i, image in enumerate(images):
        # 2. TIFF/PNG ga temporary ga save chesthunnam
        temp_img_path = f"temp_page_{i}.png"
        image.save(temp_img_path, "PNG")
        
        # 3. OCR Analysis - Ikkada Text and Layout ni identify chestundi
        # 'hocr' or 'data' use chesi location kuda track cheyochu, but simple text focus chesthunnam
        text = pytesseract.image_to_string(image)
        
        # 4. Word loki Editable Text ni add chesthunnam
        paragraph = word_doc.add_paragraph(text)
        
        # Optional: Mee request prakaram images kuda undali ante:
        # word_doc.add_picture(temp_img_path, width=Inches(6))
        
        word_doc.add_page_break()
        
        # Temporary image delete
        if os.path.exists(temp_img_path):
            os.remove(temp_img_path)

    word_doc.save(docx_path)

# ================= API ENDPOINT =================
@router.post("/pdf-to-word-ocr")
async def pdf_to_word(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Only PDF files allowed")

    file_id = str(uuid.uuid4())
    pdf_path = os.path.join(UPLOAD_DIR, f"{file_id}.pdf")
    docx_path = os.path.join(OUTPUT_DIR, f"{file_id}.docx")

    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    try:
        # Ikkada image logic run avthundi
        pdf_to_images_to_word(pdf_path, docx_path)
    except Exception as e:
        raise HTTPException(500, f"OCR Conversion failed: {str(e)}")

    return {"status": "success", "download_id": file_id}